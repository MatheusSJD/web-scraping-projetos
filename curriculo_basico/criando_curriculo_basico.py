from docx import Document
from docx.shared import Cm, Inches

curriculo = Document()

margens = curriculo.sections
for margem in margens:
    margem.top_margin = Cm(1.5)
    margem.bottom_margin = Cm(1.5)
    margem.left_margin = Cm(1.5)
    margem.right_margin = Cm(1.5)

nome = curriculo.add_heading('Matheus Santos', 0)

endereco = curriculo.add_paragraph(style="List Bullet")
endereco.add_run('Endereço: ').bold = True
endereco.add_run('Rua xxxxxxxxx').bold = False

telefone = curriculo.add_paragraph(style="List Bullet")
telefone.add_run('Telefone: ').bold = True
telefone.add_run('11920020-xxxx').bold = False

email = curriculo.add_paragraph(style="List Bullet")
email.add_run('E-mail: ').bold = True
email.add_run('mat******d9@gmail.com').bold = False

github = curriculo.add_paragraph(style="List Bullet")
github.add_run('Github: ').bold = True
github.add_run('https://github.com/matheuss0xf').bold = False

objetivo = curriculo.add_paragraph('Objetivo', style="Heading 1")
texto = curriculo.add_paragraph('\nÉ um facto estabelecido de que um leitor é distraído pelo conteúdo legível de uma página quando analisa a sua mancha gráfica. Logo, o uso de Lorem Ipsum leva a uma distribuição mais ou menos normal de letras, ao contrário do uso de "Conteúdo aqui, conteúdo aqui", tornando-o texto legível. Muitas ferramentas de publicação electrónica e editores de páginas web usam actualmente o Lorem Ipsum como o modelo de texto usado por omissão, e uma pesquisa por "lorem ipsum" irá encontrar muitos websites ainda na sua infância. Várias versões têm evoluído ao longo dos anos, por vezes por acidente, por vezes propositadamente (como no caso do humor).')

projetos = curriculo.add_paragraph('Projetos pessoais', style="Heading 1")

for projeto in range(1, 3):
    projetos = curriculo.add_paragraph(f'Projeto {projeto}', style="Heading 2")
    texto = curriculo.add_paragraph('\nExistem muitas variações das passagens do Lorem Ipsum disponíveis, mas a maior parte sofreu alterações de alguma forma, pela injecção de humor, ou de palavras aleatórias que nem sequer parecem suficientemente credíveis. Se vai usar uma passagem do Lorem Ipsum, deve ter a certeza que não contém nada de embaraçoso escondido no meio do texto.')

experiencias = curriculo.add_paragraph('Experiência', style="Heading 1")
for experiencia in range(1, 3):
    nome_empresa = curriculo.add_paragraph(f'NOME DA EMPRESA {projeto} \n', style="Heading 2")
    texto = curriculo.add_paragraph('Data: [Inicio] - [FIM]')
    texto = curriculo.add_paragraph('Atividades feitas', style="List Bullet")
    texto = curriculo.add_paragraph('Atividades feitas', style="List Bullet")
    texto = curriculo.add_paragraph('Atividades feitas',style="List Bullet")
    
formacao_academica = curriculo.add_paragraph('Experiência', style="Heading 1")
for formacao in range(1, 3):
    nome_empresa = curriculo.add_paragraph(f'NOME DA EMPRESA {projeto} \n', style="Heading 2")
    texto = curriculo.add_paragraph('Data: [Inicio] - [FIM]')
    texto = curriculo.add_paragraph('Nome do curso', style="List Bullet")
    texto = curriculo.add_paragraph('Descrição sobre o curso', style="List Bullet")
    
    
curriculo.save('curriculoDemo.docx')
